import os

os.environ['TF_ENABLE_ONEDNN_OPTS'] = '0'

import re
import json
import time
import base64
import logging
import threading
import warnings
import requests
import sqlite3
import mimetypes
import pandas as pd
import numpy as np
from datetime import datetime, timezone
from urllib.parse import urlparse
from tqdm import tqdm
from flask import (
    Flask, g, redirect, url_for, render_template, session, request, flash,
    send_file, has_request_context, copy_current_request_context, jsonify, abort
)
from langdetect import detect
from google_auth_oauthlib.flow import Flow
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from google.auth.transport.requests import Request
from bs4 import BeautifulSoup, Comment, MarkupResemblesLocatorWarning
from bleach import clean
from tensorflow.keras.models import load_model
import re
import tldextract 
from PIL import Image
import io
import PyPDF2
from docx import Document
import pickle
import re
import yara
from sklearn.model_selection import train_test_split

# NEW: import Roberta body classifier
from body_classifier import predict_body_label

# NEW: import RoBERTa document classifier for attachments
from document_classifier import initialize_document_classifier, classify_document, classify_text_content

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'nice-122423-1234')

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize RoBERTa document classifier for text attachments (replaces old spaCy model)
initialize_document_classifier()

# REMOVED: Old spaCy model for attachment classification
# nlp = spacy.load("trained_model_txt_file")

# Load CNN model for image classification
image_model = load_model("image_model.h5")

# Custom Jinja2 filter for splitting strings
def jinja_split(value, delimiter):
    return value.split(delimiter)

os.environ["CUDA_LAUNCH_BLOCKING"] = "1"

# Register the filter with the Flask app
app.jinja_env.filters['split'] = jinja_split

# YARA configuration - use relative path from project directory
app.config['YARA_RULES_DIR'] = os.path.join(os.path.dirname(__file__), 'awesome-yara', 'rules')
app.config['TEMP_DIR'] = os.path.join(os.getcwd(), 'temp')


import pandas as pd

# Load trusted domains/emails from CSV (updated)
TRUSTED_CSV_PATH = "top-1m.csv"
trusted_set = set()

# Public email providers that should NOT be domain-whitelisted
PUBLIC_EMAIL_PROVIDERS = {
    'gmail.com', 'yahoo.com', 'hotmail.com', 'outlook.com', 'live.com',
    'aol.com', 'icloud.com', 'mail.com', 'protonmail.com', 'yandex.com'
}

try:
    df = pd.read_csv(TRUSTED_CSV_PATH, header=None)
    # Filter out public providers during loading as extra safety
    trusted_set = set(str(x).strip().lower() for x in df[0].dropna() 
                     if str(x).strip().lower() not in PUBLIC_EMAIL_PROVIDERS)
    logger.info(f"Loaded {len(trusted_set)} trusted entries from {TRUSTED_CSV_PATH}")
except Exception as e:
    logger.error(f"Error loading trusted CSV: {e}")
    trusted_set = set()



warnings.filterwarnings("ignore", category=MarkupResemblesLocatorWarning)
yara_rules = None

def extract_and_classify_urls(subject, body):
    """
    Extract URLs from email subject and body, classify them as Safe or Potentially Phishing
    based on trusted_set (loaded from CSV).
    Returns a list of dictionaries: {url, domain, status}
    """
    try:
        soup = BeautifulSoup(body, 'html.parser')
        plain_text_body = soup.get_text(separator=' ', strip=True)
        text = f"{subject or ''} {plain_text_body}"

        # Regex: extract http/https/www URLs
        url_pattern = re.compile(
            r'(https?://[\w\.-]+\.\w+[\w\.:/?=&%-]*|www\.[\w\.-]+\.\w+[\w\.:/?=&%-]*)',
            re.IGNORECASE
        )

        urls = url_pattern.findall(text)
        url_list = []

        for url in urls:
            # Normalize URL prefix (add http:// if starts with www)
            if url.lower().startswith('www.'):
                normalized_url = 'http://' + url
            else:
                normalized_url = url
            if normalized_url not in url_list:
                url_list.append(normalized_url)

        url_info = []
        for url in url_list:
            try:
                parsed = urlparse(url)
                hostname = parsed.hostname

                if not hostname:
                    continue
                hostname = hostname.lower()

                # Normalize to base domain using tldextract
                ext = tldextract.extract(hostname)
                normalized_domain = f"{ext.domain}.{ext.suffix}" if ext.suffix else hostname

                # Check against trusted_set
                if normalized_domain in trusted_set:
                    status = "Safe"
                else:
                    status = "Potentially Phishing"

                url_info.append({
                    "url": url,
                    "domain": normalized_domain,
                    "status": status
                })

            except Exception as e:
                logger.warning(f"Error processing URL {url}: {e}")
                continue

        logger.debug(f"Extracted {len(url_info)} URLs: {url_info}")
        return url_info

    except Exception as e:
        logger.error(f"Error extracting URLs: {e}")
        return []


# REMOVED: get_spacy_email_prediction — replaced by Roberta module

def initialize_yara_rules():
    global yara_rules
    rules_dir = app.config['YARA_RULES_DIR']
    if not os.path.exists(rules_dir):
        logger.warning(f"YARA rules directory not found: {rules_dir}. YARA scanning will be disabled.")
        yara_rules = None
        return
    try:
        rule_files = []
        for root, dirs, files in os.walk(rules_dir):
            for file in files:
                if file.endswith(('.yar', '.yara')):
                    rule_path = os.path.join(root, file)
                    try:
                        with open(rule_path, 'r') as f:
                            content = f.read()
                            yara.compile(source=content)
                        rule_files.append((os.path.splitext(os.path.basename(file))[0], rule_path))
                    except yara.SyntaxError as e:
                        logger.warning(f"Skipping invalid YARA rule {file}: {e}")
                        continue
        if not rule_files:
            logger.warning("No valid YARA rule files found. YARA scanning will be disabled.")
            yara_rules = None
            return
        yara_rules = yara.compile(filepaths={rule_name: rule_path for rule_name, rule_path in rule_files})
        logger.info("YARA rules loaded and compiled successfully.")
    except Exception as e:
        logger.warning(f"Error loading YARA rules: {e}. YARA scanning will be disabled.")
        yara_rules = None

def scan_attachment_with_yara(attachment_data, filename):
    logger.info(f"Starting YARA scan for file: {filename}")
    try:
        if yara_rules is None:
            logger.error("YARA rules are not initialized")
            return {
                'status': 'error',
                'message': "YARA rules are not initialized",
                'details': None
            }

        temp_dir = app.config['TEMP_DIR']
        os.makedirs(temp_dir, exist_ok=True)
        temp_file_path = os.path.join(temp_dir, f"temp_{filename}")
        with open(temp_file_path, 'wb') as f:
            f.write(attachment_data)
        matches = yara_rules.match(temp_file_path)
        os.remove(temp_file_path)

        if matches:
            match_details = [f"{match.rule} ({', '.join(match.tags)})" for match in matches]
            logger.info(f"YARA scan completed for {filename}: Malicious patterns detected")
            return {
                'status': 'unsafe',
                'message': f"Malicious patterns detected: {', '.join(match_details)}",
                'details': matches
            }
        else:
            logger.info(f"YARA scan completed for {filename}: No malicious patterns detected")
            return {
                'status': 'safe',
                'message': "No malicious patterns detected",
                'details': None
            }
    except Exception as e:
        logger.error(f"Error during YARA scan for {filename}: {e}")
        return {
            'status': 'error',
            'message': f"Scan failed: {str(e)}",
            'details': None
        }

# Google Gmail API credentials from environment variables
GOOGLE_CLIENT_ID = os.environ.get('GOOGLE_CLIENT_ID')
GOOGLE_CLIENT_SECRET = os.environ.get('GOOGLE_CLIENT_SECRET')
GOOGLE_REDIRECT_URI = os.environ.get('GOOGLE_REDIRECT_URI', 'http://127.0.0.1:5000/callback')
SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'

# Build Google OAuth client config dynamically from environment variables
def get_google_client_config():
    return {
        "web": {
            "client_id": GOOGLE_CLIENT_ID,
            "client_secret": GOOGLE_CLIENT_SECRET,
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "redirect_uris": [GOOGLE_REDIRECT_URI]
        }
    }

class SenderWarningFilter(logging.Filter):
    def filter(self, record):
        if record.levelno == logging.WARNING and "Failed to extract email from sender" in record.msg:
            return False
        if record.levelno == logging.WARNING and "Sender lacks domain (no '@' symbol)" in record.msg:
            return False
        return True

logger.addFilter(SenderWarningFilter())
global_lock = threading.Lock()

def get_db_connection():
    """
    Get a database connection, either from g.db (if in a request context) or a new connection.
    Enables WAL mode for better concurrency.
    """
    if has_request_context():
        if hasattr(g, 'db') and g.db is not None:
            return g.db
    conn = sqlite3.connect('emails.db', check_same_thread=False)
    conn.execute('PRAGMA journal_mode=WAL')
    conn.row_factory = sqlite3.Row
    return conn

def safe_db_cursor():
    """
    Get a cursor for database operations, ensuring thread safety.
    """
    conn = get_db_connection()
    return conn.cursor(), conn

@app.before_request
def before_request():
    g.db = get_db_connection()

@app.teardown_request
def teardown_request(exception):
    if hasattr(g, 'db'):
        g.db.close()

import re
import tldextract

def is_trusted_email_or_domain(email: str) -> bool:
    """
    Check if email address or its domain is in trusted_set (from CSV).
    Rules:
    1. First try exact email match (e.g., ceo@bank.com)
    2. Then try domain match, BUT exclude public email providers
    3. Public providers like gmail.com require exact email match only
    """
    if not email:
        return False
    
    email = email.strip().lower()
    logger.debug(f"Checking trust for email: {email}")

    # Step 1: Exact email match (highest priority)
    if email in trusted_set:
        logger.info(f"Email {email} found in trusted set (exact match)")
        return True

    # Step 2: Extract domain part
    match = re.search(r'[\w\.-]+@([\w\.-]+\.\w+)', email)
    if not match:
        logger.debug(f"Could not extract domain from {email}")
        return False
    
    domain_raw = match.group(1).lower()
    
    # Step 3: Normalize domain using tldextract
    ext = tldextract.extract(domain_raw)
    normalized_domain = f"{ext.domain}.{ext.suffix}" if ext.suffix else ext.domain
    logger.debug(f"Normalized domain: {normalized_domain}")

    # Step 4: Check if it's a public email provider
    if normalized_domain in PUBLIC_EMAIL_PROVIDERS:
        logger.debug(f"Domain {normalized_domain} is a public provider - requires exact email match only")
        return False  # Public providers require exact email match only

    # Step 5: Check if corporate domain is trusted
    is_domain_trusted = normalized_domain in trusted_set
    if is_domain_trusted:
        logger.info(f"Email {email} trusted via corporate domain: {normalized_domain}")
    else:
        logger.debug(f"Domain {normalized_domain} not in trusted set")
    
    return is_domain_trusted


def extract_domain(sender):
    """
    Extract the domain from the sender's email address.
    """
    if not isinstance(sender, str):
        logger.warning(f"Sender is not a string: {sender}")
        return ""
    sender = sender.strip()
    if not sender or sender.lower() in ['undisclosed-recipients:;', 'undisclosed-recipients', '']:
        logger.debug(f"Sender is a placeholder: {sender}")
        return ""
    domain_match = re.search(r'@([\w\.-]+\.\w+)', sender)
    if domain_match:
        domain = domain_match.group(1).lower()
        logger.debug(f"Extracted domain: {domain} from sender: {sender}")
        typo_corrections = {
            "gamil.com": "gmail.com",
            "hotmai.com": "hotmail.com",
            "outlok.com": "outlook.com",
            "yahoomail.com": "yahoo.com"
        }
        corrected_domain = typo_corrections.get(domain, domain)
        logger.debug(f"Domain after typo correction: {corrected_domain}")
        return corrected_domain
    logger.warning(f"Could not extract domain from sender: {sender}")
    return ""


def extract_email(sender):
    """
    Extract the email address from the sender string.
    """
    if not isinstance(sender, str):
        logger.warning(f"Sender is not a string: {sender}")
        return ""
    sender = sender.strip()
    if not sender or sender.lower() in ['undisclosed-recipients:;', 'undisclosed-recipients', '']:
        logger.debug(f"Sender is a placeholder: {sender}")
        return ""
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', sender)
    if email_match:
        return email_match.group(0).lower()
    logger.warning(f"Failed to extract email from sender: {sender}")
    return ""

def classify_email(email_id, sender_email, subject, content, attachments):
    """
    Classifies an email using a multi-layered approach with asymmetric security-focused thresholds.
    - SAFE: Requires 90% confidence
    - PHISHING: Requires only 35% confidence  
    - Between 35-89%: Flagged for manual review
    Returns category, confidence, explanation [(name, raw_weight)], needs_review flag, and raw factors dict.
    """
    try:
        # Security-focused thresholds
        SAFE_THRESHOLD = 0.90      # Need 90% confidence to mark as SAFE
        PHISHING_THRESHOLD = 0.35   # Only 35% confidence needed for PHISHING
        
        logger.info(f"=== Classifying Email {email_id} from {sender_email} ===")
        
        # 1) Check CSV whitelist (exact email or corporate domain, excluding public providers)
        is_whitelisted = is_trusted_email_or_domain(sender_email)
        logger.info(f"CSV Whitelist check: {is_whitelisted}")
        
        if is_whitelisted:
            logger.info(f"Email {email_id} WHITELISTED as Safe - skipping model")
            return "Safe", 100.0, [], False, {}

        # 2) Language detection (non-English -> unidentified)
        text_for_language = f"{subject or ''} {content or ''}".strip()
        logger.debug(f"Text for classification (first 200 chars): '{text_for_language[:200]}'")
        
        if not text_for_language:
            return "Un-indentified Email", 100.0, [], True, {}
        try:
            detected_lang = detect(text_for_language)
            logger.debug(f"Detected language: {detected_lang}")
            if detected_lang != 'en':
                return "Un-indentified Email", 100.0, [], True, {}
        except Exception as e:
            logger.warning(f"Lang-detect failed for {email_id}: {e}")

        # 3) Initialize raw factor scores
        factors = {
            'ai_model_prediction': 0.0,   # phishing probability from Roberta
            'url_analysis': 0.0,
            'attachment_analysis': 0.0,
            'content_analysis': 0.0,
            'sender_trust': 0.0
        }

        # 4) AI body classifier (Roberta + LoRA → phishing probability)
        try:
            label, conf, probs = predict_body_label(text_for_language)
            logger.info(f"Roberta prediction: {label}, confidence: {conf}")
            logger.debug(f"Roberta probabilities: {probs}")
            
            # Store raw probabilities for transparency
            factors['ai_model_probs'] = probs

            # Always prefer explicit phishing probability from model if available
            if isinstance(probs, dict) and 'Phishing' in probs:
                phish_prob = float(probs['Phishing'])
            else:
                phish_prob = float(conf) if label.lower() == 'phishing' else float(1.0 - conf)

            # Clamp to [0,1] and store as phishing likelihood
            factors['ai_model_prediction'] = max(0.0, min(1.0, phish_prob))
            logger.debug(f"AI phishing factor: {factors['ai_model_prediction']}")
        except Exception as e:
            logger.error(f"Roberta body classifier failed for email {email_id}: {e}")

        # 5) URL reputation (check each URL domain against CSV set)
        urls = extract_and_classify_urls(subject, content)
        logger.debug(f"Extracted URLs: {urls}")
        if any(u['status'] == 'Potentially Phishing' for u in urls):
            factors['url_analysis'] = 0.7
            logger.debug(f"URL analysis factor: {factors['url_analysis']}")

        # 6) Attachment risk analysis
        attachment_risk = 0.0
        for att in attachments:
            yara_result = scan_attachment_with_yara(att[1], att[0])
            if yara_result.get('status') == 'unsafe':
                attachment_risk = 1.0
                logger.info(f"YARA detected malicious attachment: {att[0]}")
                break
            if att[2] == 'sensitive':  # att[2] is "sensitivity" from spacy/CNN
                attachment_risk = max(attachment_risk, 0.6)
                logger.debug(f"Sensitive attachment detected: {att[0]}")
        factors['attachment_analysis'] = attachment_risk
        logger.debug(f"Attachment factor: {factors['attachment_analysis']}")

        # 7) Content heuristics (suspicious keywords)
        content_lower = (content or '').lower()
        suspicious_keywords = ['urgent', 'verify', 'security alert', 'password', 'click here']
        found_keywords = [kw for kw in suspicious_keywords if kw in content_lower]
        if found_keywords:
            factors['content_analysis'] = 0.4
            logger.debug(f"Suspicious keywords found: {found_keywords}")

        # 8) Sender trust heuristics (suspicious TLDs)
        suspicious_tlds = ['.xyz', '.biz', '.info', '.top', '.loan']
        normalized_domain = extract_domain(sender_email)
        if any(normalized_domain.endswith(tld) for tld in suspicious_tlds):
            factors['sender_trust'] = 0.5
            logger.debug(f"Suspicious TLD detected: {normalized_domain}")

        # 9) Weighted ensemble calculation
        weights = {
            'ai_model_prediction': 0.40,    # Roberta phishing probability
            'url_analysis': 0.25,           # URL reputation
            'attachment_analysis': 0.15,     # YARA + sensitivity
            'content_analysis': 0.10,        # Keyword heuristics  
            'sender_trust': 0.10            # Domain reputation
        }
        
        weighted_score = sum(factors.get(k, 0.0) * w for k, w in weights.items())
        model_confidence = max(0.0, min(1.0, weighted_score))
        
        logger.info(f"Factor scores: {factors}")
        logger.info(f"Weighted score: {weighted_score:.4f}, Model confidence: {model_confidence:.4f}")

        # 10) Apply asymmetric security-focused thresholds
        needs_review_threshold = False
        
        if model_confidence >= SAFE_THRESHOLD:
            category = "Safe"
            logger.info(f"HIGH CONFIDENCE SAFE: {model_confidence:.4f} >= {SAFE_THRESHOLD}")
        elif model_confidence >= PHISHING_THRESHOLD:
            category = "Phishing"
            logger.info(f"PHISHING DETECTED: {model_confidence:.4f} >= {PHISHING_THRESHOLD}")
        else:
            # Very low confidence - default to safe but flag for review
            category = "Safe"
            needs_review_threshold = True
            logger.warning(f"LOW CONFIDENCE: {model_confidence:.4f} < {PHISHING_THRESHOLD} - flagged for review")

        # Convert to percentage for UI display
        confidence = round(model_confidence * 100, 2)

        # 11) Build explanation list (only non-zero contributing factors)
        mapping = {
            'ai_model_prediction': 'Roberta Body Model (Phish Prob)',
            'url_analysis': 'URL Analysis',
            'attachment_analysis': 'Attachment Analysis', 
            'content_analysis': 'Content Analysis',
            'sender_trust': 'Sender Trust'
        }
        explanation = [(mapping[k], v) for k, v in factors.items() if k in mapping and v > 0]

        # 12) Comprehensive review flagging
        no_signals = (
            factors.get('ai_model_prediction', 0.0) == 0.0 and
            factors.get('url_analysis', 0.0) == 0.0 and
            factors.get('attachment_analysis', 0.0) == 0.0 and
            factors.get('content_analysis', 0.0) == 0.0 and
            factors.get('sender_trust', 0.0) == 0.0
        )
        
        # Flag for review if: no signals, uncertain threshold zone, or very low confidence
        needs_review = no_signals or needs_review_threshold or confidence < 25
        
        logger.info(f"FINAL DECISION: {category} (confidence: {confidence}%), needs_review: {needs_review}")
        
        return category, confidence, explanation, needs_review, factors

    except Exception as e:
        logger.error(f"Classification error for email from {sender_email}: {e}", exc_info=True)
        return "Unknown", 0.0, [], True, {
            'ai_model_prediction': 0.0,
            'url_analysis': 0.0,
            'attachment_analysis': 0.0,
            'content_analysis': 0.0,
            'sender_trust': 0.0
        }


def build_gmail_service(credentials):
    return build('gmail', 'v1', credentials=credentials)

def clean_preview_text(html_content, max_length=200):
    if not html_content:
        return ""
    soup = BeautifulSoup(html_content, 'html.parser')
    for tag in soup.find_all(True):
        tag.unwrap()
    text = soup.get_text(separator=' ', strip=True)
    text = ' '.join(text.split())
    if len(text) > max_length:
        text = text[:max_length] + "..."
    return text

def classify_text_attachment(text):
    """
    Classify text content as sensitive or non-sensitive using RoBERTa document classifier.
    (Replaced old spaCy model with new RoBERTa + LoRA model)
    """
    try:
        if not text or text.strip() == '':
            return 'non-sensitive'
        # Use the new RoBERTa document classifier
        return classify_text_content(text)
    except Exception as e:
        logger.error(f"Error classifying text attachment: {e}")
        return 'non-sensitive'

def classify_image_attachment(image_data):
    """
    Classify image content as sensitive or non-sensitive using CNN.
    """
    try:
        image = Image.open(io.BytesIO(image_data))
        if image.mode == 'RGBA':
            image = image.convert('RGB')
        image = image.resize((148, 148))
        image_array = np.array(image) / 255.0
        image_array = np.expand_dims(image_array, axis=0)
        prediction = image_model.predict(image_array)
        sensitivity = 'sensitive' if prediction[0] > 0.5 else 'non-sensitive'
        return sensitivity
    except Exception as e:
        logger.error(f"Error classifying image attachment: {e}")
        return 'non-sensitive'

def fetch_and_store_emails(service, provider, access_token=None, max_results=500, label='INBOX'):
    try:
        user_email = session.get('user_email', 'unknown_user')
        logger.info(f"Fetching emails for user: {user_email}")

        if provider == 'gmail':
            results = service.users().messages().list(
                userId='me',
                labelIds=[label],
                maxResults=max_results,
                q=''
            ).execute()
            messages = results.get('messages', [])
        elif provider == 'outlook':
            folder = 'sentitems' if label == 'SENT' else 'inbox'
            headers = {'Authorization': f'Bearer {access_token}'}
            response = requests.get(
                f'https://graph.microsoft.com/v1.0/me/mailfolders/{folder}/messages',
                headers=headers,
                params={'$top': max_results, '$orderby': 'receivedDateTime desc'}
            )
            if response.status_code != 200:
                raise Exception(f"Failed to fetch emails: {response.text}")
            messages = response.json().get('value', [])

        logger.info(f"Fetched {len(messages)} {label.lower()} messages from {provider.capitalize()} for user {user_email}.")

        for msg in messages:
            try:
                cursor, conn = safe_db_cursor()
                message_id = msg['id']
                cursor.execute('SELECT category FROM Email WHERE message_id = ? AND provider = ? AND user_email = ?', (message_id, provider, user_email))
                existing_email = cursor.fetchone()
                if existing_email:
                    logger.info(f"Email {message_id} already classified as {existing_email['category']} for user {user_email}, skipping reclassification.")
                    continue

                if provider == 'gmail':
                    email_data = process_message(service, message_id)
                    if email_data is None:
                        logger.warning(f"Email data is None for message ID {message_id}, skipping.")
                        continue
                    msg_full = service.users().messages().get(userId='me', id=message_id, format='full').execute()
                    received_date = int(msg_full.get('internalDate', 0))
                    receiver = email_data.get('receiver', 'Unknown Receiver')
                elif provider == 'outlook':
                    email_data = process_outlook_email(msg)
                    if email_data is None:
                        logger.warning(f"Email data is None for message ID {msg['id']}, skipping.")
                        continue
                    received_date_str = msg.get('receivedDateTime', '1970-01-01T00:00:00Z')
                    received_date_dt = datetime.strptime(received_date_str, '%Y-%m-%dT%H:%M:%SZ')
                    received_date = int(received_date_dt.timestamp() * 1000)
                    receiver = email_data.get('receiver', 'Unknown Receiver')

                logger.info(f"Processing email: {email_data['subject']} for user {user_email}, received_date: {received_date}")
                logger.info(f"Email body: {email_data['body'][:100]}...")

                conn.execute('''
                    INSERT OR IGNORE INTO Email
                    (message_id, sender, receiver, subject, body, headers, category, provider, label, has_feedback, user_email, received_date, confidence_score, needs_review, explanation, features, urls)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 0, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    email_data['message_id'],
                    email_data['sender'],
                    receiver,
                    email_data['subject'],
                    email_data['body'],
                    email_data['headers'],
                    email_data['category'],
                    provider,
                    label,
                    user_email,
                    received_date,
                    email_data['confidence_score'],
                    1 if email_data['needs_review'] else 0,
                    json.dumps(email_data.get('explanation', [])),
                    json.dumps(email_data.get('features', {})),
                    json.dumps(email_data.get('urls', []))
                ))
                email_id = conn.execute(
                    'SELECT id FROM Email WHERE message_id = ? AND user_email = ?',
                    (email_data['message_id'], user_email)
                ).fetchone()
                if not email_id:
                    logger.error(f"Failed to retrieve email ID for message ID {message_id}")
                    conn.rollback()
                    continue
                email_id = email_id['id']

                for attachment in email_data.get('attachments', []):
                    if len(attachment) == 5:
                        filename, data, sensitivity, content_type, content_id = attachment
                        conn.execute('''
                            INSERT OR IGNORE INTO Attachment
                            (filename, data, email_id, content_id, sensitivity, content_type)
                            VALUES (?, ?, ?, ?, ?, ?)
                        ''', (filename, data, email_id, content_id, sensitivity, content_type))
                    elif len(attachment) == 4:
                        filename, data, sensitivity, content_type = attachment
                        conn.execute('''
                            INSERT OR IGNORE INTO Attachment
                            (filename, data, email_id, sensitivity, content_type)
                            VALUES (?, ?, ?, ?, ?)
                        ''', (filename, data, email_id, sensitivity, content_type))
                conn.commit()
            except Exception as e:
                logger.error(f"Error processing email {msg.get('id', 'unknown_id')} for user {user_email}: {str(e)}")
                conn.rollback()
                continue
            finally:
                if not hasattr(conn, '__enter__'):
                    conn.close()
    except Exception as e:
        logger.error(f"Error fetching {label.lower()} emails from {provider.capitalize()} for user {user_email}: {str(e)}")
        flash(f'Error fetching emails: {str(e)}', 'error')

import csv

def append_trusted_entries_to_csv(entries):
    """
    Append one or more trusted entries (emails or domains) to CSV safely.
    - 'entries' is an iterable of strings (already lowercased).
    - Skips values already present in trusted_set.
    - Updates trusted_set in-memory after writing.
    """
    global trusted_set
    to_add = [e for e in set(entries) if e and e not in trusted_set]
    if not to_add:
        return 0

    try:
        # Append mode, newline='' prevents blank lines on Windows
        with open(TRUSTED_CSV_PATH, 'a', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            for item in to_add:
                writer.writerow([item])
        # Update in-memory set so subsequent checks pick it up immediately
        trusted_set.update(to_add)
        logger.info(f"Appended {len(to_add)} trusted entries to {TRUSTED_CSV_PATH}: {to_add[:5]}{'...' if len(to_add) > 5 else ''}")
        return len(to_add)
    except Exception as e:
        logger.error(f"Failed to append trusted entries to CSV: {e}")
        return 0

def read_pdf_content(file_data):
    """Extract text from PDF files."""
    text = ''
    reader = PyPDF2.PdfReader(io.BytesIO(file_data))
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + '\n'
    return text

def read_docx_content(file_data):
    """Extract text from DOCX files."""
    doc = Document(io.BytesIO(file_data))
    return "\n".join([para.text for para in doc.paragraphs])

# Outlook API credentials from environment variables
OUTLOOK_CLIENT_ID = os.environ.get('OUTLOOK_CLIENT_ID')
OUTLOOK_CLIENT_SECRET = os.environ.get('OUTLOOK_CLIENT_SECRET')
OUTLOOK_REDIRECT_URI = os.environ.get('OUTLOOK_REDIRECT_URI', 'http://localhost:5000/callback_outlook')
OUTLOOK_SCOPES = ['https://graph.microsoft.com/.default']

@app.route('/authorize_outlook')
def authorize_outlook():
    auth_url = f"https://login.microsoftonline.com/common/oauth2/v2.0/authorize?"
    params = {
        'client_id': OUTLOOK_CLIENT_ID,
        'response_type': 'code',
        'redirect_uri': OUTLOOK_REDIRECT_URI,
        'response_mode': 'query',
        'scope': ' '.join(OUTLOOK_SCOPES),
        'state': os.urandom(16).hex()
    }
    session['state'] = params['state']
    return redirect(auth_url + '&'.join([f"{k}={v}" for k, v in params.items()]))

@app.route('/callback_outlook')
def callback_outlook():
    code = request.args.get('code')
    if not code:
        return 'Authorization failed', 400

    num_emails = session.get('num_emails', 10)
    data = {
        'grant_type': 'authorization_code',
        'code': code,
        'redirect_uri': OUTLOOK_REDIRECT_URI,
        'client_id': OUTLOOK_CLIENT_ID,
        'client_secret': OUTLOOK_CLIENT_SECRET
    }
    token_response = requests.post('https://login.microsoftonline.com/common/oauth2/v2.0/token', data=data)
    if token_response.status_code != 200:
        error_data = token_response.json()
        error_description = error_data.get('error_description', 'Unknown error')
        logger.error(f"Token request failed: {error_description}")
        if 'invalid_grant' in error_data.get('error', ''):
            return redirect(url_for('authorize_outlook'))
        return f"Error getting token: {error_description}", 500

    access_token = token_response.json().get('access_token')
    headers = {'Authorization': f'Bearer {access_token}'}
    user_response = requests.get('https://graph.microsoft.com/v1.0/me', headers=headers)
    if user_response.status_code == 200:
        user_email = user_response.json().get('userPrincipalName', 'unknown_user')
        logger.info(f"Logged in Outlook user: {user_email}")
        session['user_email'] = user_email
    else:
        logger.error(f"Failed to fetch user email: {user_response.text}")
        session['user_email'] = 'unknown_user'

    fetch_and_store_emails(None, 'outlook', access_token=access_token, max_results=num_emails)
    return redirect(url_for('outlook_dashboard'))

@app.route('/scan_attachment/<int:attachment_id>')
def scan_attachment(attachment_id):
    cursor = g.db.cursor()
    cursor.execute('SELECT data, filename FROM Attachment WHERE id = ?', (attachment_id,))
    attachment = cursor.fetchone()
    if not attachment:
        return jsonify({'status': 'error', 'message': 'Attachment not found'}), 404
    filename = attachment['filename']
    attachment_data = attachment['data']
    scan_result = scan_attachment_with_yara(attachment_data, filename)
    return jsonify(scan_result)

@app.route('/set_email_count', methods=['POST'])
def set_email_count():
    num_emails = int(request.form.get('num_emails', 10))
    session['num_emails'] = num_emails
    return redirect(url_for('index'))

@app.route('/outlook_dashboard')
def outlook_dashboard():
    user_email = session.get('user_email', 'unknown_user')
    cursor = g.db.cursor()
    cursor.execute('SELECT * FROM Email WHERE provider = "outlook" AND label = "INBOX" AND user_email = ? ORDER BY COALESCE(received_date, 0) DESC', (user_email,))
    emails = cursor.fetchall()
    cursor.execute('SELECT * FROM Attachment WHERE email_id IN (SELECT id FROM Email WHERE user_email = ?)', (user_email,))
    attachments = cursor.fetchall()
    return render_template(
        'outlook_dashboard.html',
        emails=emails,
        all_attachments=attachments,
        clean_preview_text=clean_preview_text
    )

@app.route('/submit_feedback/<int:email_id>', methods=['POST'])
def submit_feedback(email_id):
    user_email = session.get('user_email', 'unknown_user')
    cursor = g.db.cursor()
    cursor.execute('SELECT * FROM Email WHERE id = ? AND user_email = ?', (email_id, user_email))
    email_row = cursor.fetchone()
    if not email_row:
        flash('Email not found or you do not have access to this email!', 'error')
        return redirect(url_for('email_details', email_id=email_id))

    email = dict(email_row)
    correct_category = request.form.get('correct_category')
    feedback_reason = request.form.get('feedback_reason', '').strip()

    if correct_category not in ['Safe', 'Phishing']:
        flash('Invalid category selected!', 'error')
        return redirect(url_for('email_details', email_id=email_id))

    # CSV-based trust persistence
    if correct_category == 'Safe':
        # Extract exact email and normalized base domain
        sender_email_str = email.get('sender', '')
        exact_email = extract_email(sender_email_str)  # e.g., noreply@google.com
        # Normalize domain robustly using tldextract
        dom_raw = extract_domain(sender_email_str)     # might be sub.mail.google.co.uk
        if dom_raw:
            ext = tldextract.extract(dom_raw.lower())
            base_domain = f"{ext.domain}.{ext.suffix}" if ext.suffix else ext.domain
        else:
            base_domain = ''

        # Decide what to store:
        # - Always prefer storing the base domain to benefit future users broadly
        # - If you also want to whitelist the exact address (when manually added), add both
        to_store = []
        if base_domain:
            to_store.append(base_domain.lower())
        if exact_email:
            # Optional: only add exact email if you want both address and domain whitelisting
            # Comment this line out if you want domain-only persistence
            to_store.append(exact_email.lower())

        added = append_trusted_entries_to_csv(to_store)
        if added > 0:
            logger.info(f"Feedback whitelisted entries added: {to_store}")
        else:
            logger.info(f"Feedback entries already present or nothing to add: {to_store}")

    # Persist feedback metadata on the Email row
    cursor.execute('''
        UPDATE Email
        SET category = ?, has_feedback = 1, feedback_reason = ?
        WHERE id = ? AND user_email = ?
    ''', (correct_category, feedback_reason, email_id, user_email))
    g.db.commit()

    flash('Feedback submitted successfully!', 'success')
    return redirect(url_for('email_details', email_id=email_id, category=correct_category))


@app.route('/outlook_email/<int:email_id>')
def outlook_email_details(email_id):
    cursor = g.db.cursor()
    cursor.execute('SELECT * FROM Email WHERE id = ? AND provider = "outlook"', (email_id,))
    email_row = cursor.fetchone()
    if not email_row:
        flash('Email not found!', 'error')
        return redirect(url_for('outlook_dashboard'))

    email = dict(email_row)
    try:
        email['explanation'] = json.loads(email['explanation']) if email['explanation'] else []
        email['features'] = json.loads(email['features']) if email['features'] else {}
        email['urls'] = json.loads(email['urls']) if email['urls'] else []
    except json.JSONDecodeError as e:
        logger.error(f"Error parsing JSON fields for email ID {email_id}: {e}")
        email['explanation'] = []
        email['features'] = {}
        email['urls'] = []

    if isinstance(email['explanation'], list):
        email['explanation'] = [(item[0], item[1]) for item in email['explanation']]
    else:
        email['explanation'] = []

    cursor.execute('SELECT id, filename, sensitivity, content_type FROM Attachment WHERE email_id = ?', (email_id,))
    attachments = [dict(row) for row in cursor.fetchall()]

    email_body = email['body'] if email['body'] is not None else ''
    soup = BeautifulSoup(email_body, 'html.parser')
    for img in soup.find_all('img'):
        src = img.get('src')
        if src and src.lower().startswith('cid:'):
            content_id = src.split(':')[1].strip().lower()
            img['src'] = url_for('serve_image', email_id=email_id, content_id=content_id)
    modified_body = soup.get_text(separator=' ', strip=True)

    return render_template('email_details.html',
                           email=email,
                           attachments=attachments,
                           modified_body=modified_body)

def urlize(text):
    """
    Convert URLs in text to clickable HTML links.
    """
    if not text:
        return text
    url_pattern = r'(https?://[^\s]+|www\.[^\s]+)'
    return re.sub(url_pattern, lambda match: f'{match.group(0)}' if match.group(0).startswith('www.') else f'{match.group(0)}', text)

app.jinja_env.filters['urlize'] = urlize

def extract_text_from_pdf(pdf_data):
    """Extract text from PDF data."""
    text = ''
    reader = PyPDF2.PdfReader(io.BytesIO(pdf_data))
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + '\n'
    return text

def extract_text_from_docx(docx_data):
    """Extract text from DOCX data."""
    doc = Document(io.BytesIO(docx_data))
    return "\n".join([para.text for para in doc.paragraphs])

def process_message(service, message_id):
    """
    Processes a raw Gmail message, recursively extracting the body, attachments,
    and other metadata, then classifies it.
    """
    try:
        msg = service.users().messages().get(userId='me', id=message_id, format='full').execute()
        headers = {h['name']: h['value'] for h in msg['payload']['headers']}
        sender = headers.get('From', 'Unknown Sender')
        if not re.search(r'[\w\.-]+@[\w\.-]+\.\w+', sender):
            sender = 'unknown@unknown.com'
        receiver = headers.get('To', 'Unknown Receiver')
        subject = headers.get('Subject', 'No Subject')

        attachments = []
        body = ""
        html_content = None
        plain_content = None

        # DFS parts
        parts_to_process = [msg['payload']]
        while parts_to_process:
            part = parts_to_process.pop(0)
            mime_type = part.get('mimeType', '')

            if 'parts' in part:
                parts_to_process = part['parts'] + parts_to_process
                continue

            if mime_type == 'text/html':
                if 'data' in part['body']:
                    html_content = base64.urlsafe_b64decode(part['body']['data']).decode('utf-8', errors='ignore')
            elif mime_type == 'text/plain':
                if 'data' in part['body'] and not html_content:
                    plain_content = base64.urlsafe_b64decode(part['body']['data']).decode('utf-8', errors='ignore')

            # Attachments
            if part.get('filename'):
                if part['body'].get('attachmentId'):
                    att_id = part['body']['attachmentId']
                    att = service.users().messages().attachments().get(
                        userId='me',
                        messageId=message_id,
                        id=att_id
                    ).execute()
                    data = base64.urlsafe_b64decode(att['data'])
                    file_type = part.get('mimeType')
                    sensitivity = 'non-sensitive'
                    if file_type.startswith('image/'):
                        sensitivity = classify_image_attachment(data)
                    elif file_type == 'application/pdf':
                        text = extract_text_from_pdf(data)
                        sensitivity = classify_text_attachment(text)
                    elif file_type == 'application/msword':
                        text = extract_text_from_docx(data)
                        sensitivity = classify_text_attachment(text)
                    elif file_type == 'text/plain':
                        text = data.decode('utf-8', errors='ignore')
                        sensitivity = classify_text_attachment(text)

                    attachments.append((
                        part['filename'],
                        data,
                        sensitivity,
                        file_type,
                        part.get('headers', [{}])[0].get('value', '').strip('<>')
                    ))

        # Fallback for simple emails
        if not (html_content or plain_content) and msg['payload'].get('body', {}).get('data'):
            body_data = msg['payload']['body']['data']
            plain_content = base64.urlsafe_b64decode(body_data).decode('utf-8', errors='ignore')

        # Clean body
        if html_content:
            soup = BeautifulSoup(html_content, 'html.parser')
            for tag in soup.find_all(['script', 'style', 'link', 'meta']):
                tag.decompose()
            for comment in soup.find_all(string=lambda text: isinstance(text, Comment)):
                comment.extract()
            body = soup.get_text(separator=' ', strip=True)
        elif plain_content:
            body = plain_content
        body = ' '.join(body.split())
        if body.startswith(subject):
            body = body[len(subject):].strip()

        urls = extract_and_classify_urls(subject, body)
        category, confidence, explanation, needs_review, features = classify_email(message_id, sender, subject, body, attachments)

        return {
            'message_id': message_id,
            'sender': sender,
            'receiver': receiver,
            'subject': subject,
            'body': body,
            'headers': json.dumps(headers),
            'category': category,
            'confidence_score': confidence,
            'explanation': explanation,
            'needs_review': needs_review,
            'features': features,
            'attachments': attachments,
            'urls': urls
        }
    except Exception as e:
        logger.error(f"Error processing message {message_id}: {str(e)}")
        return {
            'message_id': message_id,
            'sender': 'Unknown Sender',
            'receiver': 'Unknown Receiver',
            'subject': 'No Subject',
            'body': '',
            'headers': json.dumps({}),
            'category': 'Safe',
            'confidence_score': 0.0,
            'explanation': [],
            'needs_review': True,
            'features': {},
            'attachments': [],
            'urls': []
        }

def process_outlook_email(email):
    try:
        sender = email['from']['emailAddress']['address'] if 'from' in email and 'emailAddress' in email['from'] else 'Unknown Sender'
        if not re.search(r'[\w\.-]+@[\w\.-]+\.\w+', sender):
            sender = 'unknown@unknown.com'
        receiver = email['toRecipients'][0]['emailAddress']['address'] if 'toRecipients' in email and email['toRecipients'] else 'Unknown Receiver'
        date = email.get('receivedDateTime', None)
        subject = email.get('subject', 'No Subject')
        body = email.get('body', {}).get('content', '')

        soup = BeautifulSoup(body, 'html.parser')
        for tag in soup.find_all(['style', 'script']):
            tag.decompose()
        for comment in soup.find_all(string=lambda text: isinstance(text, Comment)):
            comment.extract()
        body = soup.get_text(separator=' ', strip=True)
        body = ' '.join(body.split())
        if body.startswith(subject):
            body = body[len(subject):].strip()

        urls = extract_and_classify_urls(subject, body)

        # Initialize attachments list
        attachments = []
        if 'attachments' in email:
            for attachment in email['attachments']:
                if attachment['contentType'].startswith('image/'):
                    sensitivity = classify_image_attachment(base64.b64decode(attachment['contentBytes']))
                elif attachment['contentType'] == 'application/pdf':
                    text = extract_text_from_pdf(base64.b64decode(attachment['contentBytes']))
                    sensitivity = classify_text_attachment(text)
                elif attachment['contentType'] == 'application/msword':
                    text = extract_text_from_docx(base64.b64decode(attachment['contentBytes']))
                    sensitivity = classify_text_attachment(text)
                elif attachment['contentType'] == 'text/plain':
                    text = base64.b64decode(attachment['contentBytes']).decode('utf-8', errors='ignore')
                    sensitivity = classify_text_attachment(text)
                else:
                    sensitivity = 'non-sensitive'
                attachments.append((
                    attachment['name'],
                    base64.b64decode(attachment['contentBytes']),
                    sensitivity,
                    attachment['contentType'],
                    attachment.get('contentId', ''),
                ))

        category, confidence, explanation, needs_review, features = classify_email(email['id'], sender, subject, body, attachments)

        return {
            'message_id': email['id'],
            'sender': sender,
            'receiver': receiver,
            'subject': subject,
            'body': body,
            'headers': json.dumps({}),
            'category': category,
            'confidence_score': confidence,
            'explanation': explanation,
            'needs_review': needs_review,
            'features': features,
            'attachments': attachments,
            'urls': urls
        }
    except Exception as e:
        logger.error(f"Error processing Outlook email {email.get('id', 'unknown_id')}: {str(e)}")
        return {
            'message_id': email.get('id', 'unknown_id'),
            'sender': 'Unknown Sender',
            'receiver': 'Unknown Receiver',
            'subject': 'No Subject',
            'body': '',
            'headers': json.dumps({}),
            'category': 'Unknown',
            'confidence_score': 0.0,
            'explanation': [],
            'needs_review': True,
            'features': {},
            'attachments': [],
            'urls': []
        }

@app.route('/sent_dashboard')
def sent_dashboard():
    user_email = session.get('user_email', 'unknown_user')
    logger.info(f"Displaying sent dashboard for user: {user_email}")
    cursor = g.db.cursor()
    cursor.execute('SELECT * FROM Email WHERE provider = "gmail" AND label = "SENT" AND user_email = ? ORDER BY COALESCE(received_date, 0) DESC', (user_email,))
    emails = cursor.fetchall()
    cursor.execute('SELECT * FROM Attachment WHERE email_id IN (SELECT id FROM Email WHERE user_email = ?)', (user_email,))
    attachments = cursor.fetchall()
    return render_template(
        'dashboard.html',
        emails=emails,
        all_attachments=attachments,
        clean_preview_text=clean_preview_text
    )

@app.route('/outlook_sent_dashboard')
def outlook_sent_dashboard():
    user_email = session.get('user_email', 'unknown_user')
    cursor = g.db.cursor()
    cursor.execute('SELECT * FROM Email WHERE provider = "outlook" AND label = "SENT" AND user_email = ? ORDER BY COALESCE(received_date, 0) DESC', (user_email,))
    emails = cursor.fetchall()
    cursor.execute('SELECT * FROM Attachment WHERE email_id IN (SELECT id FROM Email WHERE user_email = ?)', (user_email,))
    attachments = cursor.fetchall()
    return render_template(
        'outlook_dashboard.html',
        emails=emails,
        all_attachments=attachments,
        clean_preview_text=clean_preview_text
    )

@app.route('/email/<int:email_id>')
def email_details(email_id):
    user_email = session.get('user_email', 'unknown_user')
    cursor = g.db.cursor()
    cursor.execute('SELECT * FROM Email WHERE id = ? AND user_email = ?', (email_id, user_email))
    email_row = cursor.fetchone()
    if not email_row:
        flash('Email not found or you do not have access to this email!', 'error')
        return redirect(url_for('dashboard'))

    email = dict(email_row)
    try:
        email['explanation'] = json.loads(email['explanation']) if email['explanation'] else []
        email['features'] = json.loads(email['features']) if email['features'] else {}
        email['urls'] = json.loads(email['urls']) if email['urls'] else []
    except json.JSONDecodeError as e:
        logger.error(f"Error parsing JSON fields for email ID {email_id}: {e}")
        email['explanation'] = []
        email['features'] = {}
        email['urls'] = []

    if isinstance(email['explanation'], list):
        email['explanation'] = [(item[0], item[1]) for item in email['explanation']]
    else:
        email['explanation'] = []

    cursor.execute('SELECT id, filename, sensitivity, content_type FROM Attachment WHERE email_id = ?', (email_id,))
    attachments = [dict(row) for row in cursor.fetchall()]

    soup = BeautifulSoup(email['body'], 'html.parser')
    for tag in soup.find_all(['style', 'script']):
        tag.decompose()
    for comment in soup.find_all(string=lambda text: isinstance(text, Comment)):
        comment.extract()
    modified_body = soup.get_text(separator=' ', strip=True)

    return render_template('email_details.html',
                           email=email,
                           attachments=attachments,
                           modified_body=modified_body)

@app.route('/authorize_gmail')
def authorize_gmail():
    flow = Flow.from_client_config(
        get_google_client_config(),
        scopes=SCOPES,
        redirect_uri=url_for('callback', _external=True, _scheme='http')
    )
    authorization_url, state = flow.authorization_url(
        access_type='offline',
        include_granted_scopes='true',
        prompt='consent'
    )
    session['state'] = state
    return redirect(authorization_url)

@app.route('/callback')
def callback():
    num_emails = session.get('num_emails', 10)
    logger.info(f"Fetching {num_emails} emails")

    flow = Flow.from_client_config(
        get_google_client_config(),
        scopes=SCOPES,
        state=session['state'],
        redirect_uri=url_for('callback', _external=True, _scheme='http')
    )
    flow.fetch_token(authorization_response=request.url)
    credentials = flow.credentials
    service = build_gmail_service(credentials)
    try:
        profile = service.users().getProfile(userId='me').execute()
        user_email = profile.get('emailAddress', 'unknown_user')
        logger.info(f"Logged in user: {user_email}")
        session['user_email'] = user_email
    except Exception as e:
        logger.error(f"Error fetching user email: {e}")
        session['user_email'] = 'unknown_user'

    cursor = g.db.cursor()
    cursor.execute('DELETE FROM Email WHERE user_email != ? OR user_email IS NULL', (user_email,))
    g.db.commit()
    logger.info(f"Cleared emails for users other than {user_email}")

    fetch_and_store_emails(service, 'gmail', label='INBOX', max_results=num_emails)
    fetch_and_store_emails(service, 'gmail', label='SENT', max_results=num_emails)
    return redirect(url_for('dashboard'))

@app.route('/dashboard')
def dashboard():
    user_email = session.get('user_email', 'unknown_user')
    logger.info(f"Displaying dashboard for user: {user_email}")
    cursor = g.db.cursor()
    cursor.execute('SELECT * FROM Email WHERE provider = "gmail" AND label = "INBOX" AND user_email = ? ORDER BY COALESCE(received_date, 0) DESC', (user_email,))
    emails = cursor.fetchall()
    logger.info(f"Dashboard emails for {user_email}: {len(emails)} emails found")
    cursor.execute('SELECT * FROM Attachment WHERE email_id IN (SELECT id FROM Email WHERE user_email = ?)', (user_email,))
    attachments = cursor.fetchall()
    return render_template(
        'dashboard.html',
        emails=emails,
        all_attachments=attachments,
        clean_preview_text=clean_preview_text
    )

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/debug_emails')
def debug_emails():
    cursor = g.db.cursor()
    cursor.execute('SELECT * FROM Email')
    emails = cursor.fetchall()
    return render_template('debug_emails.html', emails=emails)

@app.route('/attachment/<int:attachment_id>')
def serve_attachment(attachment_id):
    cursor = g.db.cursor()
    cursor.execute('''
        SELECT data, filename, sensitivity, content_type
        FROM Attachment
        WHERE id = ?
    ''', (attachment_id,))
    attachment = cursor.fetchone()
    if not attachment:
        return "Attachment not found", 404

    filename = attachment['filename']
    attachment_data = attachment['data']
    sensitivity = attachment['sensitivity']
    content_type = attachment['content_type']

    if sensitivity == 'sensitive':
        flash("Warning: This attachment is classified as sensitive. Proceed with caution.", "warning")
        scan_result = scan_attachment_with_yara(attachment_data, filename)
        if scan_result['status'] == 'unsafe':
            logger.warning(f"Malicious attachment detected: {filename} - {scan_result['message']}")
            flash(f"Warning: Attachment '{filename}' contains suspicious content. {scan_result['message']}", "danger")
            return redirect(url_for('dashboard'))
        if scan_result['status'] == 'error':
            logger.error(f"YARA scan error for attachment {filename}: {scan_result['message']}")
            flash(f"Error scanning attachment: {scan_result['message']}", "warning")

    return send_file(
        io.BytesIO(attachment_data),
        mimetype=content_type,
        download_name=filename,
        as_attachment=True
    )

@app.route('/image/<int:email_id>/<content_id>')
def serve_image(email_id, content_id):
    cursor = g.db.cursor()
    cursor.execute('''
        SELECT data, content_type
        FROM Attachment
        WHERE email_id = ? AND content_id = ?
    ''', (email_id, content_id.lower()))
    attachment = cursor.fetchone()
    if attachment:
        return send_file(
            io.BytesIO(attachment['data']),
            mimetype=attachment['content_type'],
            as_attachment=False
        )
    return "Image not found", 404

if __name__ == '__main__':
    with app.app_context():
        db = get_db_connection()
        cursor = db.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS Email (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                message_id TEXT UNIQUE,
                sender TEXT,
                receiver TEXT,
                subject TEXT,
                body TEXT,
                headers TEXT,
                category TEXT,
                provider TEXT,
                label TEXT,
                has_feedback INTEGER DEFAULT 0,
                user_email TEXT,
                received_date INTEGER,
                confidence_score REAL,
                needs_review INTEGER DEFAULT 0,
                feedback_reason TEXT,
                explanation TEXT,
                features TEXT,
                urls TEXT
            )
        ''')
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS Attachment (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                filename TEXT,
                data BLOB,
                email_id INTEGER,
                content_id TEXT,
                sensitivity TEXT,
                content_type TEXT,
                FOREIGN KEY (email_id) REFERENCES Email (id)
            )
        ''')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_user_email ON Email (user_email)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_message_id ON Email (message_id)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_category ON Email (category)')
        db.commit()
        db.close()

    initialize_yara_rules()
    app.run(debug=True, use_reloader=False)
